import os
from typing import List, Dict, Any
import fitz  # PyMuPDF
import docx

from app.utils.logger import logger


class DocumentParserService:
    """Service to parse and extract text content from various document formats (PDF, DOCX, TXT)."""

    def parse_pages(self, file_path: str, extension: str) -> List[Dict[str, Any]]:
        """Parse clean text from a file segmented by page numbers.

        Args:
            file_path (str): Absolute or relative path to the file.
            extension (str): File extension with or without leading dot.

        Returns:
            List[Dict[str, Any]]: List of dictionary items containing 'page_number' and 'text'.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If file content is empty or file format is unsupported.
            RuntimeError: If document parsing encounters a library error.
        """
        logger.info(f"Parsing Started: {file_path} (Format: {extension})")
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found at: {file_path}")

        ext = extension.lower().lstrip(".")
        if ext == "pdf":
            return self._parse_pdf(file_path)
        elif ext == "docx":
            return self._parse_docx(file_path)
        elif ext in ("txt", "text"):
            return self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")

    def parse(self, file_path: str, extension: str) -> str:
        """Parse and extract the entire text from a file as a single concatenated string.

        Args:
            file_path (str): Path to the file.
            extension (str): File extension.

        Returns:
            str: Full concatenated clean text.
        """
        pages = self.parse_pages(file_path, extension)
        return "\n".join(page["text"] for page in pages).strip()

    def _parse_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        try:
            doc = fitz.open(file_path)
            pages = []
            for i, page in enumerate(doc):
                text = page.get_text().strip()
                # Include page even if empty to preserve page count structure
                pages.append({
                    "page_number": i + 1,
                    "text": text
                })
            doc.close()
            
            # Check if all pages are completely empty
            if not any(page["text"] for page in pages):
                raise ValueError("PDF document is empty or contains no extractable text.")
                
            return pages
        except ValueError as ve:
            logger.error(f"Validation failure during PDF parsing: {ve}")
            raise ve
        except Exception as e:
            logger.error(f"PyMuPDF failed to parse PDF at {file_path}: {e}", exc_info=True)
            raise RuntimeError(f"PDF parsing failure: {e}")

    def _parse_docx(self, file_path: str) -> List[Dict[str, Any]]:
        try:
            doc = docx.Document(file_path)
            text_parts = []
            
            # Extract paragraphs text
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())

            # Extract tables text
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        text_parts.append(" | ".join(row_cells))

            full_text = "\n".join(text_parts).strip()
            if not full_text:
                raise ValueError("DOCX document is empty or contains no extractable text.")

            # DOCX has no explicit native pages in python-docx, treat entire doc as page 1
            return [{
                "page_number": 1,
                "text": full_text
            }]
        except ValueError as ve:
            logger.error(f"Validation failure during DOCX parsing: {ve}")
            raise ve
        except Exception as e:
            logger.error(f"python-docx failed to parse DOCX at {file_path}: {e}", exc_info=True)
            raise RuntimeError(f"DOCX parsing failure: {e}")

    def _parse_txt(self, file_path: str) -> List[Dict[str, Any]]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read().strip()
            if not content:
                raise ValueError("TXT document is empty.")
            
            # TXT document has no pages, treat as page 1
            return [{
                "page_number": 1,
                "text": content
            }]
        except ValueError as ve:
            logger.error(f"Validation failure during TXT parsing: {ve}")
            raise ve
        except Exception as e:
            logger.error(f"Failed to read TXT file at {file_path}: {e}", exc_info=True)
            raise RuntimeError(f"TXT parsing failure: {e}")


# Singleton service instance
parser_service = DocumentParserService()
